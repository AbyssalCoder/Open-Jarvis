"""
Tool functions that the Brain agent can dispatch.
Each tool takes string args and returns a string result.
"""

import json
import os
import subprocess
import httpx
import base64
import smtplib
import platform
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from pathlib import Path
from datetime import datetime, timedelta


async def web_search(query: str) -> str:
    """Search the web using DuckDuckGo instant answer API."""
    try:
        async with httpx.AsyncClient(timeout=10.0) as client:
            # DuckDuckGo Instant Answer API (no key needed)
            resp = await client.get(
                "https://api.duckduckgo.com/",
                params={"q": query, "format": "json", "no_html": "1", "skip_disambig": "1"},
            )
            data = resp.json()

            results = []
            # Abstract (main answer)
            if data.get("Abstract"):
                results.append(f"**{data.get('Heading', 'Result')}**: {data['Abstract']}")
                if data.get("AbstractURL"):
                    results.append(f"Source: {data['AbstractURL']}")

            # Answer (instant answer)
            if data.get("Answer"):
                results.append(f"Answer: {data['Answer']}")

            # Related topics
            for topic in data.get("RelatedTopics", [])[:5]:
                if isinstance(topic, dict) and topic.get("Text"):
                    results.append(f"- {topic['Text'][:200]}")

            if results:
                return "\n".join(results)
            else:
                return f"No instant results found for '{query}'. Try asking me to explain it instead."
    except Exception as e:
        return f"Search error: {e}"


async def get_weather(city: str = "auto") -> str:
    """Get current weather using wttr.in."""
    try:
        url = f"https://wttr.in/{city}?format=j1" if city != "auto" else "https://wttr.in/?format=j1"
        async with httpx.AsyncClient(timeout=10.0) as client:
            resp = await client.get(url)
            data = resp.json()
            current = data.get("current_condition", [{}])[0]
            area = data.get("nearest_area", [{}])[0]
            city_name = area.get("areaName", [{}])[0].get("value", "Unknown")
            country = area.get("country", [{}])[0].get("value", "")
            temp_c = current.get("temp_C", "?")
            desc = current.get("weatherDesc", [{}])[0].get("value", "")
            humidity = current.get("humidity", "?")
            feels = current.get("FeelsLikeC", "?")
            return f"Weather in {city_name}, {country}: {desc}, {temp_c}°C (feels {feels}°C), humidity {humidity}%"
    except Exception as e:
        return f"Weather error: {e}"


async def run_terminal_command(command: str) -> str:
    """Execute a terminal command and return output. CAUTION: sandboxed."""
    # Safety: block obviously dangerous commands
    blocked = ["rm -rf /", "format ", "del /s /q", "rmdir /s", ":(){", "fork bomb"]
    cmd_lower = command.lower()
    for b in blocked:
        if b in cmd_lower:
            return f"Blocked dangerous command: {command}"

    try:
        result = subprocess.run(
            command, shell=True, capture_output=True, text=True, timeout=30,
            cwd=str(Path.home()),
        )
        output = result.stdout[:2000] if result.stdout else ""
        error = result.stderr[:500] if result.stderr else ""
        if error and not output:
            return f"Error:\n{error}"
        return output or "(no output)"
    except subprocess.TimeoutExpired:
        return "Command timed out (30s limit)"
    except Exception as e:
        return f"Command error: {e}"


async def capture_screen() -> str:
    """Take a screenshot and return base64 data."""
    try:
        import mss
        with mss.mss() as sct:
            monitor = sct.monitors[1]  # Primary monitor
            img = sct.grab(monitor)
            # Convert to PNG bytes
            from mss.tools import to_png
            png_bytes = to_png(img.rgb, img.size)
            b64 = base64.b64encode(png_bytes).decode()
            return f"Screenshot captured ({img.size.width}x{img.size.height}). Base64 length: {len(b64)}"
    except ImportError:
        return "Screen capture requires 'mss' package. Install with: pip install mss"
    except Exception as e:
        return f"Screenshot error: {e}"


async def get_system_info() -> str:
    """Get current system information."""
    import platform
    import psutil
    try:
        cpu = psutil.cpu_percent(interval=0.5)
        mem = psutil.virtual_memory()
        disk = psutil.disk_usage("/")
        return (
            f"System: {platform.system()} {platform.release()}\n"
            f"CPU: {cpu}% used ({psutil.cpu_count()} cores)\n"
            f"RAM: {mem.percent}% used ({mem.used // (1024**3)}GB / {mem.total // (1024**3)}GB)\n"
            f"Disk: {disk.percent}% used ({disk.used // (1024**3)}GB / {disk.total // (1024**3)}GB)\n"
            f"Time: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
        )
    except ImportError:
        return "System info requires 'psutil' package."
    except Exception as e:
        return f"System info error: {e}"


async def read_file(filepath: str) -> str:
    """Read a file and return its contents (truncated to 3000 chars)."""
    try:
        p = Path(filepath).expanduser().resolve()
        if not p.exists():
            return f"File not found: {filepath}"
        if p.stat().st_size > 1_000_000:
            return f"File too large (>{1_000_000} bytes): {filepath}"
        content = p.read_text(encoding="utf-8", errors="replace")
        if len(content) > 3000:
            return content[:3000] + f"\n... (truncated, {len(content)} total chars)"
        return content
    except Exception as e:
        return f"Read error: {e}"


async def write_file(filepath: str, content: str) -> str:
    """Write content to a file."""
    try:
        p = Path(filepath).expanduser().resolve()
        p.parent.mkdir(parents=True, exist_ok=True)
        p.write_text(content, encoding="utf-8")
        return f"Written {len(content)} chars to {p}"
    except Exception as e:
        return f"Write error: {e}"


async def get_datetime() -> str:
    """Get current date, time, and day."""
    now = datetime.now()
    return now.strftime("Date: %A, %B %d, %Y\nTime: %I:%M:%S %p\nTimezone: Local")


# ─── Vision / Image Analysis ───────────────────────────────────────────────

# Global YOLO model cache (load once, reuse)
_yolo_model = None


def _get_yolo_model():
    """Load YOLOv8 model (cached singleton). Uses GPU if available, else CPU."""
    global _yolo_model
    if _yolo_model is None:
        try:
            from ultralytics import YOLO
            import torch

            _yolo_model = YOLO("yolov8n.pt")  # nano model: fast, 6MB

            # Move to GPU if CUDA is available — offloads RAM to VRAM
            if torch.cuda.is_available():
                _yolo_model.to("cuda")
            elif hasattr(torch, "backends") and hasattr(torch.backends, "mps") and torch.backends.mps.is_available():
                _yolo_model.to("mps")  # Apple Silicon GPU
        except Exception:
            pass
    return _yolo_model


async def _yolo_detect(image_bytes: bytes) -> str | None:
    """Run YOLO object detection on image bytes. Returns description of detected objects."""
    import asyncio
    try:
        model = _get_yolo_model()
        if model is None:
            return None

        # Save to temp file for YOLO (it accepts file paths)
        import tempfile
        with tempfile.NamedTemporaryFile(suffix=".jpg", delete=False) as f:
            f.write(image_bytes)
            temp_path = f.name

        # Run inference in thread pool to not block async loop
        loop = asyncio.get_event_loop()
        results = await loop.run_in_executor(None, lambda: model(temp_path, verbose=False))

        # Clean up temp file
        try:
            os.unlink(temp_path)
        except OSError:
            pass

        if not results or len(results) == 0:
            return "No objects detected in the image."

        result = results[0]
        detections = []
        if result.boxes is not None and len(result.boxes) > 0:
            # Count objects by class
            class_counts = {}
            for box in result.boxes:
                cls_id = int(box.cls[0])
                cls_name = result.names[cls_id]
                conf = float(box.conf[0])
                if conf >= 0.3:  # confidence threshold
                    class_counts[cls_name] = class_counts.get(cls_name, 0) + 1

            if class_counts:
                for name, count in sorted(class_counts.items(), key=lambda x: -x[1]):
                    if count > 1:
                        detections.append(f"{count} {name}s")
                    else:
                        detections.append(f"1 {name}")

        if detections:
            return f"Detected objects: {', '.join(detections)}"
        return "Image captured but no recognizable objects detected."

    except ImportError:
        return None
    except Exception as e:
        return f"YOLO detection error: {e}"


async def _gemini_vision(b64_image: str, query: str, mime_type: str = "image/png") -> str | None:
    """Analyze an image using Ollama vision model (local) with YOLO detections as context."""
    # We no longer use Gemini — everything is local
    return None


async def _ollama_describe(yolo_result: str, query: str) -> str:
    """Ask Ollama to describe a scene based on YOLO detections + user query."""
    from config import config
    ollama_url = config.ollama_url

    prompt = (
        f"You are JARVIS, an AI vision assistant. "
        f"A camera captured an image and object detection found: {yolo_result}\n\n"
        f"The user asked: \"{query}\"\n\n"
        f"Based on the detected objects, give a natural, helpful description of what you see. "
        f"Be conversational and specific. If you can infer context from the objects (e.g., person at a desk with laptop = working), do so."
    )

    try:
        async with httpx.AsyncClient(timeout=30.0) as client:
            resp = await client.post(
                f"{ollama_url}/api/generate",
                json={
                    "model": "jarvis:latest",
                    "prompt": prompt,
                    "stream": False,
                    "options": {"num_predict": 300, "temperature": 0.7, "num_gpu": -1, "num_ctx": 2048},
                },
            )
            if resp.status_code == 200:
                return resp.json().get("response", "").strip()
    except Exception:
        pass
    return None


async def vision_analyze(query: str = "Describe what you see") -> str:
    """Capture the screen and analyze with multi-model pipeline (YOLO + GroundingDINO + LLM)."""
    try:
        import mss
        from vision.pipeline import get_vision_pipeline

        # Capture screen
        with mss.mss() as sct:
            monitor = sct.monitors[1]
            img = sct.grab(monitor)
            from mss.tools import to_png
            png_bytes = to_png(img.rgb, img.size)

        # Use the multi-model pipeline
        pipeline = get_vision_pipeline()
        result = await pipeline.analyze(png_bytes, query)
        return result.get("result", "No analysis returned")

    except ImportError:
        return "Vision requires 'mss' package. Install with: pip install mss"
    except Exception as e:
        return f"Vision analysis error: {e}"


# ─── Productivity Scheduler ────────────────────────────────────────────────

REMINDERS_FILE = Path(os.environ.get("APPDATA", str(Path.home()))) / "JARVIS" / "reminders.json"


def _load_reminders() -> list[dict]:
    """Load reminders from disk."""
    try:
        if REMINDERS_FILE.exists():
            return json.loads(REMINDERS_FILE.read_text(encoding="utf-8"))
    except Exception:
        pass
    return []


def _save_reminders(reminders: list[dict]):
    """Persist reminders to disk."""
    REMINDERS_FILE.parent.mkdir(parents=True, exist_ok=True)
    REMINDERS_FILE.write_text(json.dumps(reminders, indent=2, default=str), encoding="utf-8")


async def scheduler(action: str, text: str = "", time: str = "") -> str:
    """Manage reminders and schedules.
    Actions: add, list, remove, clear.
    Time formats: '5m' (minutes), '2h' (hours), '2026-05-21 14:30', 'tomorrow 9am'.
    """
    reminders = _load_reminders()

    if action == "list":
        if not reminders:
            return "No reminders set. Your schedule is clear."
        lines = []
        now = datetime.now()
        for i, r in enumerate(reminders):
            due = datetime.fromisoformat(r["due"])
            status = "OVERDUE" if due < now else f"due {due.strftime('%b %d %I:%M %p')}"
            lines.append(f"{i + 1}. {r['text']} — {status}")
        return "Your reminders:\n" + "\n".join(lines)

    elif action == "add":
        if not text:
            return "Please specify what to remind you about."
        # Parse time
        due = _parse_time(time)
        reminder = {"text": text, "due": due.isoformat(), "created": datetime.now().isoformat()}
        reminders.append(reminder)
        _save_reminders(reminders)
        return f"Reminder set: '{text}' — due {due.strftime('%A, %b %d at %I:%M %p')}"

    elif action == "remove":
        if not text:
            return "Specify the reminder number to remove (use 'list' first)."
        try:
            idx = int(text) - 1
            if 0 <= idx < len(reminders):
                removed = reminders.pop(idx)
                _save_reminders(reminders)
                return f"Removed reminder: '{removed['text']}'"
            return f"Invalid reminder number. You have {len(reminders)} reminders."
        except ValueError:
            # Try to find by text match
            matches = [r for r in reminders if text.lower() in r["text"].lower()]
            if matches:
                reminders.remove(matches[0])
                _save_reminders(reminders)
                return f"Removed reminder: '{matches[0]['text']}'"
            return f"No reminder matching '{text}' found."

    elif action == "clear":
        count = len(reminders)
        _save_reminders([])
        return f"Cleared all {count} reminders."

    return f"Unknown scheduler action: {action}. Use: add, list, remove, clear."


def _parse_time(time_str: str) -> datetime:
    """Parse flexible time strings into datetime."""
    now = datetime.now()
    if not time_str:
        return now + timedelta(hours=1)  # Default: 1 hour from now

    t = time_str.strip().lower()

    # Relative: 5m, 30m, 2h, 1d
    if t.endswith("m") and t[:-1].isdigit():
        return now + timedelta(minutes=int(t[:-1]))
    if t.endswith("h") and t[:-1].isdigit():
        return now + timedelta(hours=int(t[:-1]))
    if t.endswith("d") and t[:-1].isdigit():
        return now + timedelta(days=int(t[:-1]))

    # "tomorrow" keyword
    if "tomorrow" in t:
        tomorrow = now + timedelta(days=1)
        # Try to extract time from "tomorrow 9am" etc.
        parts = t.replace("tomorrow", "").strip()
        if parts:
            try:
                hour = int("".join(filter(str.isdigit, parts)))
                if "pm" in parts and hour < 12:
                    hour += 12
                return tomorrow.replace(hour=hour, minute=0, second=0, microsecond=0)
            except (ValueError, TypeError):
                pass
        return tomorrow.replace(hour=9, minute=0, second=0, microsecond=0)

    # Try ISO format: "2026-05-21 14:30"
    for fmt in ["%Y-%m-%d %H:%M", "%Y-%m-%d %H:%M:%S", "%m/%d/%Y %H:%M", "%d-%m-%Y %H:%M"]:
        try:
            return datetime.strptime(t, fmt)
        except ValueError:
            continue

    # Default: 1 hour from now
    return now + timedelta(hours=1)


# ─── Media Control ─────────────────────────────────────────────────────────

async def media_control(action: str) -> str:
    """Control system media playback.
    Actions: play_pause, next, previous, volume_up, volume_down, mute.
    """
    if platform.system() != "Windows":
        return "Media control currently only supported on Windows."

    try:
        import ctypes
        user32 = ctypes.windll.user32  # type: ignore[attr-defined]

        # Virtual key codes for media keys
        VK_MAP = {
            "play_pause": 0xB3,    # VK_MEDIA_PLAY_PAUSE
            "play": 0xB3,
            "pause": 0xB3,
            "next": 0xB0,         # VK_MEDIA_NEXT_TRACK
            "skip": 0xB0,
            "previous": 0xB1,    # VK_MEDIA_PREV_TRACK
            "prev": 0xB1,
            "volume_up": 0xAF,   # VK_VOLUME_UP
            "volume_down": 0xAE, # VK_VOLUME_DOWN
            "mute": 0xAD,        # VK_VOLUME_MUTE
        }

        vk = VK_MAP.get(action.lower().strip())
        if not vk:
            return f"Unknown media action: {action}. Available: {', '.join(VK_MAP.keys())}"

        # Simulate key press and release
        KEYEVENTF_EXTENDEDKEY = 0x0001
        KEYEVENTF_KEYUP = 0x0002
        user32.keybd_event(vk, 0, KEYEVENTF_EXTENDEDKEY, 0)
        user32.keybd_event(vk, 0, KEYEVENTF_EXTENDEDKEY | KEYEVENTF_KEYUP, 0)

        action_names = {
            0xB3: "Toggled play/pause",
            0xB0: "Skipped to next track",
            0xB1: "Went to previous track",
            0xAF: "Volume increased",
            0xAE: "Volume decreased",
            0xAD: "Toggled mute",
        }
        return action_names.get(vk, f"Executed {action}")

    except Exception as e:
        return f"Media control error: {e}"


# ─── Mobile Sync / Notifications ──────────────────────────────────────────

OWNER_PROFILE = {
    "name": "Aniket",
    "email": "aniketsupermails2005@gmail.com",
    "phone": "+91 7980458591",
    "instagram": "@_alive_phoenix_",
    "github": "AbyssalCoder",
}


async def mobile_sync(action: str, message: str = "", to: str = "") -> str:
    """Send notifications to your phone/email.
    Actions: send_email, profile, send_reminder, whatsapp.
    """
    if action == "profile":
        return (
            f"Owner: {OWNER_PROFILE['name']}\n"
            f"Email: {OWNER_PROFILE['email']}\n"
            f"Phone: {OWNER_PROFILE['phone']}\n"
            f"Instagram: {OWNER_PROFILE['instagram']}\n"
            f"GitHub: {OWNER_PROFILE['github']}"
        )

    elif action == "whatsapp":
        # Open WhatsApp Web with pre-filled message
        phone = to or OWNER_PROFILE["phone"].replace("+", "").replace(" ", "")
        encoded_msg = message.replace(" ", "%20") if message else ""
        url = f"https://wa.me/{phone}?text={encoded_msg}" if encoded_msg else f"https://wa.me/{phone}"
        try:
            if platform.system() == "Windows":
                os.startfile(url)  # type: ignore[attr-defined]
            else:
                subprocess.Popen(["xdg-open", url])
            return f"Opening WhatsApp chat with {to or OWNER_PROFILE['phone']}"
        except Exception as e:
            return f"WhatsApp error: {e}"

    elif action == "send_email":
        recipient = to or OWNER_PROFILE["email"]
        if not message:
            return "Please provide a message to send."

        # Check for SMTP credentials in env
        smtp_user = os.environ.get("JARVIS_EMAIL_USER", OWNER_PROFILE["email"])
        smtp_pass = os.environ.get("JARVIS_EMAIL_PASSWORD", "")

        if not smtp_pass:
            return (
                "Email sending requires your email app password. "
                "Set it in your .env file as JARVIS_EMAIL_PASSWORD=your_app_password. "
                "For Gmail, create an App Password at myaccount.google.com under Security."
            )

        try:
            msg = MIMEMultipart()
            msg["From"] = smtp_user
            msg["To"] = recipient
            msg["Subject"] = "JARVIS Notification"
            msg.attach(MIMEText(message, "plain"))

            with smtplib.SMTP_SSL("smtp.gmail.com", 465, timeout=10) as server:
                server.login(smtp_user, smtp_pass)
                server.send_message(msg)

            return f"Email sent to {recipient} successfully."
        except smtplib.SMTPAuthenticationError:
            return "Email auth failed. Make sure JARVIS_EMAIL_PASSWORD in .env is a valid Gmail App Password."
        except Exception as e:
            return f"Email sending error: {e}"

    elif action == "send_reminder":
        return await mobile_sync("send_email", message=f"JARVIS Reminder: {message}", to=OWNER_PROFILE["email"])

    return f"Unknown mobile_sync action: {action}. Available: send_email, whatsapp, profile, send_reminder."


# ─── Browser / App Opening ─────────────────────────────────────────────────

async def open_url(url: str) -> str:
    """Open a URL in the default browser."""
    try:
        import webbrowser
        webbrowser.open(url)
        return f"Opened {url} in default browser"
    except Exception as e:
        return f"Failed to open URL: {e}"


async def open_app(name: str) -> str:
    """Open an application or website by name.
    Supports: youtube, amazon, flipkart, spotify, instagram, github, google, etc.
    For YouTube videos/songs: pass full search or video URL.
    """
    name_lower = name.lower().strip()

    # URL mapping for common apps/sites
    app_urls = {
        "youtube": "https://youtube.com",
        "amazon": "https://amazon.in",
        "flipkart": "https://flipkart.com",
        "spotify": "https://open.spotify.com",
        "instagram": "https://instagram.com",
        "github": "https://github.com",
        "google": "https://google.com",
        "whatsapp": "https://web.whatsapp.com",
        "gmail": "https://mail.google.com",
        "twitter": "https://twitter.com",
        "x": "https://x.com",
        "linkedin": "https://linkedin.com",
        "reddit": "https://reddit.com",
        "chatgpt": "https://chat.openai.com",
        "amazon music": "https://music.amazon.in",
        # Shopping & delivery
        "blinkit": "https://blinkit.com",
        "zepto": "https://www.zeptonow.com",
        "bigbasket": "https://www.bigbasket.com",
        "myntra": "https://www.myntra.com",
        "meesho": "https://www.meesho.com",
        "ajio": "https://www.ajio.com",
        # Food ordering
        "zomato": "https://www.zomato.com",
        "swiggy": "https://www.swiggy.com",
    }

    # Check if it's a direct URL
    if name_lower.startswith("http://") or name_lower.startswith("https://"):
        return await open_url(name)

    # Check known apps
    for key, url in app_urls.items():
        if key in name_lower:
            return await open_url(url)

    # YouTube search / play
    if "play" in name_lower or "video" in name_lower or "song" in name_lower or "music" in name_lower:
        query = name_lower.replace("play", "").replace("video", "").replace("song", "").replace("music", "").replace("on youtube", "").strip()
        if query:
            search_url = f"https://www.youtube.com/results?search_query={query.replace(' ', '+')}"
            return await open_url(search_url)

    # Fallback: Google search to find the app/site
    search_url = f"https://www.google.com/search?q={name.replace(' ', '+')}"
    return await open_url(search_url)


async def youtube_play(query: str) -> str:
    """Search and play a video/song on YouTube — auto-plays the first result."""
    try:
        import webbrowser
        import re as _re

        # Method 1: Scrape YouTube search page for first video ID
        try:
            async with httpx.AsyncClient(timeout=10.0, follow_redirects=True) as client:
                resp = await client.get(
                    f"https://www.youtube.com/results?search_query={query.replace(' ', '+')}",
                    headers={
                        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36",
                        "Accept-Language": "en-US,en;q=0.9",
                    }
                )
                if resp.status_code == 200:
                    # Try multiple patterns to find video IDs
                    patterns = [
                        r'"videoId":"([a-zA-Z0-9_-]{11})"',
                        r'/watch\?v=([a-zA-Z0-9_-]{11})',
                        r'"url":"/watch\?v=([a-zA-Z0-9_-]{11})"',
                    ]
                    for pattern in patterns:
                        matches = _re.findall(pattern, resp.text)
                        if matches:
                            video_id = matches[0]
                            # Add autoplay=1 to force autoplay
                            video_url = f"https://www.youtube.com/watch?v={video_id}&autoplay=1"
                            webbrowser.open(video_url)
                            return f"Playing '{query}' on YouTube"
        except Exception:
            pass

        # Method 2: Use YouTube's oEmbed API to get video info (no scraping needed)
        try:
            async with httpx.AsyncClient(timeout=8.0) as client:
                # Search via Invidious API (public YouTube frontend)
                resp = await client.get(
                    f"https://vid.puffyan.us/api/v1/search?q={query.replace(' ', '+')}&type=video",
                    headers={"Accept": "application/json"}
                )
                if resp.status_code == 200:
                    results = resp.json()
                    if results and len(results) > 0:
                        video_id = results[0].get("videoId", "")
                        if video_id:
                            video_url = f"https://www.youtube.com/watch?v={video_id}&autoplay=1"
                            webbrowser.open(video_url)
                            return f"Playing '{query}' on YouTube"
        except Exception:
            pass

        # Fallback: open YouTube search page (user can click the first result)
        search_url = f"https://www.youtube.com/results?search_query={query.replace(' ', '+')}"
        webbrowser.open(search_url)
        return f"Opened YouTube search for '{query}'. Click the first result to play."
    except Exception as e:
        return f"YouTube error: {e}"


# ─── File & Folder Search ──────────────────────────────────────────────────

async def search_files(query: str, path: str = "") -> str:
    """Search for files and folders on the device.
    Query can be a filename, extension, or pattern.
    Uses fuzzy matching to find files even with approximate/misspelled names.
    Path is optional starting directory (defaults to user home + common locations).
    """
    from difflib import SequenceMatcher

    search_dirs = []
    if path:
        search_dirs = [Path(path).expanduser().resolve()]
    else:
        home = Path.home()
        search_dirs = [
            home / "Desktop",
            home / "Documents",
            home / "Downloads",
            home / "OneDrive" / "Desktop",
            home / "OneDrive" / "Documents",
            home / "OneDrive" / "Downloads",
            home / "Pictures",
            home / "Videos",
            home / "Music",
        ]

    results = []
    fuzzy_results = []  # (score, path) for fuzzy matches
    query_lower = query.lower()
    max_results = 20
    max_depth = 5  # Don't recurse too deep

    # Split query into words for partial matching
    query_words = query_lower.split()

    def _fuzzy_score(name: str) -> float:
        """Calculate fuzzy match score between query and filename."""
        name_lower = name.lower()
        # Exact substring match = highest priority
        if query_lower in name_lower:
            return 1.0
        # All query words appear in filename
        if all(w in name_lower for w in query_words):
            return 0.95
        # SequenceMatcher ratio
        ratio = SequenceMatcher(None, query_lower, name_lower).ratio()
        # Also check against stem (without extension)
        stem = Path(name).stem.lower()
        stem_ratio = SequenceMatcher(None, query_lower, stem).ratio()
        return max(ratio, stem_ratio)

    def _search_dir(directory: Path, depth: int = 0):
        """Search directory with depth limit and fuzzy matching."""
        if depth > max_depth or len(results) >= max_results:
            return
        try:
            for item in directory.iterdir():
                if len(results) >= max_results:
                    return
                try:
                    name = item.name
                    # Exact substring match
                    if query_lower in name.lower():
                        results.append(str(item))
                    else:
                        # Fuzzy match
                        score = _fuzzy_score(name)
                        if score >= 0.55:  # Threshold for fuzzy matches
                            fuzzy_results.append((score, str(item)))

                    if item.is_dir() and not item.name.startswith('.') and item.name not in (
                        'node_modules', '__pycache__', '.git', '.venv', 'venv',
                        'AppData', 'Application Data', 'Local Settings',
                    ):
                        _search_dir(item, depth + 1)
                except (PermissionError, OSError):
                    continue
        except (PermissionError, OSError):
            pass

    for search_dir in search_dirs:
        if not search_dir.exists():
            continue
        if "*" in query or "?" in query:
            try:
                for match in search_dir.rglob(query):
                    results.append(str(match))
                    if len(results) >= max_results:
                        break
            except (PermissionError, OSError):
                continue
        else:
            _search_dir(search_dir)
        if len(results) >= max_results:
            break

    # If no exact matches, use top fuzzy results
    if not results and fuzzy_results:
        fuzzy_results.sort(key=lambda x: x[0], reverse=True)
        results = [path for _, path in fuzzy_results[:15]]

    # Also add some fuzzy results to exact matches if we have room
    elif results and fuzzy_results and len(results) < max_results:
        fuzzy_results.sort(key=lambda x: x[0], reverse=True)
        for score, path in fuzzy_results[:5]:
            if path not in results and len(results) < max_results:
                results.append(path)

    if not results:
        return f"No files matching '{query}' found. Try a different name or check the spelling. I searched Desktop, Documents, Downloads, Pictures, Videos, and Music folders."

    # Format results
    is_fuzzy = not any(query_lower in Path(r).name.lower() for r in results[:3])
    header = f"Found {len(results)} result(s) for '{query}'"
    if is_fuzzy:
        header += " (closest matches)"
    lines = [header + ":"]
    for r in results[:15]:
        p = Path(r)
        kind = "folder" if p.is_dir() else f"file ({p.suffix})"
        size = ""
        if p.is_file():
            try:
                sz = p.stat().st_size
                if sz > 1_000_000:
                    size = f" [{sz // 1_000_000} MB]"
                elif sz > 1000:
                    size = f" [{sz // 1000} KB]"
            except OSError:
                pass
        lines.append(f"  {kind}: {r}{size}")

    if len(results) > 15:
        lines.append(f"  ... and {len(results) - 15} more")

    return "\n".join(lines)


# ─── Launch Desktop Applications ──────────────────────────────────────────

async def launch_app(name: str) -> str:
    """Find and launch a desktop application installed on Windows.
    Searches Start Menu shortcuts, Program Files, and common locations.
    """
    if platform.system() != "Windows":
        return "Desktop app launching currently only supported on Windows."

    name_lower = name.lower().strip()

    # 1. Try direct 'start' command for common Windows apps
    direct_apps = {
        "notepad": "notepad",
        "calculator": "calc",
        "calc": "calc",
        "paint": "mspaint",
        "cmd": "cmd",
        "powershell": "powershell",
        "terminal": "wt",
        "task manager": "taskmgr",
        "control panel": "control",
        "settings": "ms-settings:",
        "file explorer": "explorer",
        "explorer": "explorer",
        "files": "explorer",
        "word": "winword",
        "excel": "excel",
        "powerpoint": "powerpnt",
        "outlook": "outlook",
        "teams": "msteams:",
        "discord": "discord",
        "chrome": "chrome",
        "firefox": "firefox",
        "edge": "msedge",
        "brave": "brave",
        "vscode": "code",
        "vs code": "code",
        "visual studio code": "code",
        "steam": "steam",
        "vlc": "vlc",
        "obs": "obs64",
        "telegram": "telegram",
        "signal": "signal",
        "zoom": "zoom",
        "slack": "slack",
        "notion": "notion",
        "spotify": "spotify",
        "whatsapp": "whatsapp",
    }

    for key, cmd in direct_apps.items():
        if key in name_lower:
            try:
                if ":" in cmd:
                    # URI protocol (ms-settings:, msteams:, etc.)
                    os.startfile(cmd)  # type: ignore[attr-defined]
                else:
                    subprocess.Popen(cmd, shell=True, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
                return f"Launched {name}"
            except Exception:
                pass

    # 2. Search Start Menu shortcuts
    start_menu_dirs = [
        Path(os.environ.get("APPDATA", "")) / "Microsoft" / "Windows" / "Start Menu" / "Programs",
        Path(os.environ.get("PROGRAMDATA", "C:/ProgramData")) / "Microsoft" / "Windows" / "Start Menu" / "Programs",
    ]

    for menu_dir in start_menu_dirs:
        if not menu_dir.exists():
            continue
        try:
            for shortcut in menu_dir.rglob("*.lnk"):
                if name_lower in shortcut.stem.lower():
                    os.startfile(str(shortcut))  # type: ignore[attr-defined]
                    return f"Launched {shortcut.stem}"
        except Exception:
            continue

    # 3. Search Program Files for exe
    program_dirs = [
        Path("C:/Program Files"),
        Path("C:/Program Files (x86)"),
        Path(os.environ.get("LOCALAPPDATA", "")),
    ]

    for prog_dir in program_dirs:
        if not prog_dir.exists():
            continue
        try:
            for exe_file in prog_dir.rglob("*.exe"):
                if name_lower in exe_file.stem.lower() and "uninstall" not in exe_file.stem.lower():
                    subprocess.Popen(str(exe_file), shell=False, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
                    return f"Launched {exe_file.stem} from {exe_file.parent}"
        except (PermissionError, OSError):
            continue

    return f"Could not find application '{name}' on this device. Try the exact name of the app."


# ─── Browser Automation (Playwright) ─────────────────────────────────────────

def _ensure_playwright_browsers():
    """Ensure Playwright can find Chromium even when running from PyInstaller exe."""
    import os
    # Default Playwright browser location
    default_path = os.path.join(os.environ.get("LOCALAPPDATA", ""), "ms-playwright")
    if os.path.isdir(default_path):
        os.environ.setdefault("PLAYWRIGHT_BROWSERS_PATH", default_path)


# Keep Playwright references alive so browsers don't close on function return
_pw_refs = []


async def browser_action(action: str = "open", query: str = "", site: str = "google", quantity: int = 1) -> str:
    """General-purpose browser automation. Can search, open sites, add to cart, fill forms, etc.
    Actions: open, search, add_to_cart, search_product, play_video, browse, research
    Sites: amazon, flipkart, blinkit, youtube, google, or any URL
    """
    try:
        _ensure_playwright_browsers()
        from playwright.async_api import async_playwright
    except ImportError:
        return "Browser automation needs Playwright. Install with: pip install playwright && python -m playwright install chromium"

    site_lower = site.lower().strip()
    action_lower = action.lower().strip()

    try:
        # Launch browser without context manager so it stays open
        pw = await async_playwright().start()
        browser = await pw.chromium.launch(headless=False, args=["--start-maximized"])
        context = await browser.new_context(viewport={"width": 1366, "height": 900})
        page = await context.new_page()
        _pw_refs.append(pw)  # prevent garbage collection — browser stays open

        # ── Add to Cart flows ──
        if action_lower in ("add_to_cart", "search_product"):
            if site_lower == "amazon":
                return await _amazon_add_to_cart(page, query, quantity)
            elif site_lower == "flipkart":
                return await _flipkart_add_to_cart(page, query)
            elif site_lower in ("blinkit", "zepto", "bigbasket"):
                return await _grocery_search(page, query, site_lower)
            else:
                await page.goto(f"https://www.google.com/search?q={query.replace(' ', '+')}+site:{site_lower}.com")
                await page.wait_for_load_state("domcontentloaded")
                return f"Searched for '{query}' on {site}. Browser is open for you."

        # ── Play video ──
        elif action_lower == "play_video":
            return await _youtube_autoplay(page, query)

        # ── Research / Study ──
        elif action_lower == "research":
            return await _research_browse(page, query)

        # ── Find PDF ──
        elif action_lower == "find_pdf":
            await page.goto(f"https://www.google.com/search?q={query.replace(' ', '+')}+filetype:pdf", timeout=15000)
            await page.wait_for_load_state("domcontentloaded")
            await page.wait_for_timeout(2000)
            pdf_link = page.locator("a[href$='.pdf']").first
            if await pdf_link.count() > 0:
                await pdf_link.click()
                await page.wait_for_timeout(2000)
                return f"Found and opened a PDF about '{query}'. Check your browser!"
            return f"Searched for PDFs about '{query}'. Browse the results in the browser."

        # ── Open a specific site or search ──
        elif action_lower in ("open", "browse", "search"):
            if query.startswith("http://") or query.startswith("https://"):
                await page.goto(query, timeout=15000)
            elif site_lower and site_lower != "google":
                site_urls = {
                    "amazon": "https://www.amazon.in",
                    "flipkart": "https://www.flipkart.com",
                    "youtube": "https://www.youtube.com",
                    "blinkit": "https://blinkit.com",
                    "zepto": "https://www.zeptonow.com",
                    "google": "https://www.google.com",
                    "github": "https://github.com",
                    "stackoverflow": "https://stackoverflow.com",
                    "wikipedia": "https://en.wikipedia.org",
                }
                url = site_urls.get(site_lower, f"https://www.{site_lower}.com")
                await page.goto(url, timeout=15000)
                if query:
                    await page.wait_for_load_state("domcontentloaded")
                    await page.wait_for_timeout(1500)
                    search_selectors = [
                        "input[type='search']", "input[name='q']", "input[name='query']",
                        "input[name='search']", "#twotabsearchtextbox", "input[placeholder*='Search']",
                        "input[aria-label*='Search']", "input[type='text']"
                    ]
                    for sel in search_selectors:
                        try:
                            box = page.locator(sel).first
                            if await box.is_visible(timeout=1000):
                                await box.fill(query)
                                await box.press("Enter")
                                break
                        except Exception:
                            continue
            else:
                await page.goto(f"https://www.google.com/search?q={query.replace(' ', '+')}", timeout=15000)

            await page.wait_for_load_state("domcontentloaded")
            await page.wait_for_timeout(1500)
            return f"Opened browser for '{query or site}'. It's ready for you."

        else:
            await page.goto(f"https://www.google.com/search?q={query.replace(' ', '+')}", timeout=15000)
            await page.wait_for_load_state("domcontentloaded")
            return f"Searched for '{query}' in the browser."

    except Exception as e:
        error_msg = str(e)
        if "Executable doesn't exist" in error_msg:
            return "Chromium browser not found. Run this in terminal: python -m playwright install chromium"
        return f"Browser automation error: {e}"


async def _research_browse(page, query: str) -> str:
    """Open multiple research sources for a topic."""
    try:
        # Open Google Scholar first
        await page.goto(f"https://scholar.google.com/scholar?q={query.replace(' ', '+')}", timeout=15000)
        await page.wait_for_load_state("domcontentloaded")
        await page.wait_for_timeout(1500)
        return f"Opened Google Scholar with research results for '{query}'. Browse the papers and articles in your browser."
    except Exception as e:
        return f"Research browse error: {e}. Trying regular search..."


async def _amazon_add_to_cart(page, query: str, quantity: int = 1) -> str:
    """Search Amazon and add first result to cart with specified quantity."""
    try:
        await page.goto("https://www.amazon.in", timeout=30000)
        await page.wait_for_load_state("networkidle", timeout=15000)
        await page.wait_for_timeout(2000)

        # Search for product
        search_box = page.locator("#twotabsearchtextbox")
        if not await search_box.is_visible(timeout=5000):
            search_box = page.locator("input[name='field-keywords']").first
        await search_box.fill(query)
        await search_box.press("Enter")
        await page.wait_for_load_state("networkidle", timeout=15000)
        await page.wait_for_timeout(3000)

        # Click first non-sponsored product result
        results = page.locator('[data-component-type="s-search-result"]')
        result_count = await results.count()
        clicked = False
        for i in range(min(result_count, 5)):
            result = results.nth(i)
            # Skip sponsored results
            sponsored = result.locator("span:has-text('Sponsored')")
            if await sponsored.count() > 0:
                continue
            link = result.locator("h2 a").first
            if await link.count() > 0:
                await link.click()
                clicked = True
                break

        if not clicked and result_count > 0:
            link = results.first.locator("h2 a").first
            await link.click()

        # Handle new tab if opened (Amazon often opens product in new tab)
        await page.wait_for_timeout(4000)
        all_pages = page.context.pages
        if len(all_pages) > 1:
            page = all_pages[-1]
            await page.bring_to_front()
        await page.wait_for_load_state("networkidle", timeout=20000)
        await page.wait_for_timeout(2000)

        # Dismiss any popup dialogs (location, language, etc.)
        try:
            dismiss_selectors = [
                "#sp-cc-accept",
                "input[data-action-type='DISMISS']",
                "button[data-action-type='DISMISS']",
                "#a-popover-ok-button",
                ".a-popover-footer button",
            ]
            for sel in dismiss_selectors:
                btn = page.locator(sel).first
                if await btn.is_visible(timeout=500):
                    await btn.click()
                    await page.wait_for_timeout(500)
        except Exception:
            pass

        # Set quantity if > 1
        if quantity > 1:
            try:
                qty_select = page.locator("#quantity")
                if await qty_select.is_visible(timeout=3000):
                    await qty_select.select_option(str(min(quantity, 10)))
                    await page.wait_for_timeout(1000)
                else:
                    qty_input = page.locator("input[name='quantity']").first
                    if await qty_input.is_visible(timeout=1000):
                        await qty_input.fill(str(min(quantity, 10)))
                        await page.wait_for_timeout(500)
            except Exception:
                pass

        # Scroll to make Add to Cart visible
        await page.evaluate("window.scrollBy(0, 400)")
        await page.wait_for_timeout(1500)

        # Try to click "Add to Cart" — comprehensive selectors
        add_selectors = [
            "#add-to-cart-button",
            "input#add-to-cart-button",
            "#add-to-cart-button-ubb",
            "input[name='submit.add-to-cart']",
            "span#submit\\.add-to-cart > input",
            "input[value='Add to Cart']",
            "input[value='Add to cart']",
            "#addToCart input[type='submit']",
            "#addToCart button",
            "[data-csa-c-action='addToCart']",
        ]
        for sel in add_selectors:
            try:
                add_btn = page.locator(sel).first
                if await add_btn.is_visible(timeout=2000):
                    await add_btn.click(force=True)
                    await page.wait_for_timeout(3000)

                    # Check if we got to the cart confirmation page
                    cart_confirm = page.locator("#NATC_SMART_WAGON_CONF_MSG_SUCCESS, #sw-atc-confirmation, h1:has-text('Added to Cart'), #huc-v2-order-row-confirm-text")
                    if await cart_confirm.count() > 0:
                        qty_text = f" (qty: {quantity})" if quantity > 1 else ""
                        return f"Successfully added '{query}'{qty_text} to your Amazon cart! Check the browser to confirm."

                    qty_text = f" (qty: {quantity})" if quantity > 1 else ""
                    return f"Clicked 'Add to Cart' for '{query}'{qty_text}. Check the browser to confirm it was added."
            except Exception:
                continue

        # Try "Buy Now" as fallback info
        buy_btn = page.locator("#buy-now-button")
        if await buy_btn.is_visible(timeout=1000):
            return f"Found '{query}' on Amazon. 'Add to Cart' not visible but 'Buy Now' is. Page is open for you."
        return f"Found '{query}' on Amazon but couldn't find the Add to Cart button. The product page is open for you to add manually."

    except Exception as e:
        return f"Amazon: Searched for '{query}' but hit an issue ({e}). Page is open for you to continue."


async def _flipkart_add_to_cart(page, query: str) -> str:
    """Search Flipkart and add first result to cart."""
    try:
        await page.goto("https://www.flipkart.com", timeout=20000)
        await page.wait_for_load_state("networkidle", timeout=15000)
        await page.wait_for_timeout(2000)

        # Close login popup if it appears (Flipkart changes class names often)
        try:
            close_selectors = [
                "button._2KpZ6l._2doB4z",
                "button[class*='_2KpZ6l']",
                "span[role='button']:has-text('✕')",
                "button:has-text('✕')",
                "[class*='close-button']",
            ]
            for sel in close_selectors:
                btn = page.locator(sel).first
                if await btn.is_visible(timeout=2000):
                    await btn.click()
                    break
        except Exception:
            pass

        await page.wait_for_timeout(500)

        # Search for product
        search_box = page.locator("input[name='q'], input[type='text'][title='Search for products']").first
        await search_box.fill(query)
        await search_box.press("Enter")
        await page.wait_for_load_state("networkidle", timeout=15000)
        await page.wait_for_timeout(3000)

        # Click first product
        first_product = page.locator("a[href*='/p/']").first
        await first_product.click()
        await page.wait_for_timeout(3000)

        # Switch to new tab if opened
        pages = page.context.pages
        if len(pages) > 1:
            page = pages[-1]
            await page.wait_for_load_state("networkidle", timeout=15000)
        await page.wait_for_timeout(2000)

        # Try to click "ADD TO CART" (case-insensitive with multiple selectors)
        cart_selectors = [
            "button:has-text('ADD TO CART')",
            "button:has-text('Add to Cart')",
            "button:has-text('Add to cart')",
            "button:has-text('ADD TO BAG')",
            "[class*='add-to-cart']",
            "button[class*='_2AkmmA']",
        ]
        for sel in cart_selectors:
            btn = page.locator(sel).first
            try:
                if await btn.is_visible(timeout=2000):
                    await btn.click()
                    await page.wait_for_timeout(3000)
                    return f"Added '{query}' to your Flipkart cart! Check your browser to proceed."
            except Exception:
                continue

        # Try "BUY NOW" as info
        buy_btn = page.locator("button:has-text('BUY NOW'), button:has-text('Buy Now')").first
        try:
            if await buy_btn.is_visible(timeout=2000):
                return f"Found '{query}' on Flipkart. 'Buy Now' is available. Page is open for you."
        except Exception:
            pass

        return f"Found '{query}' on Flipkart. Page is open for you to add to cart."

    except Exception as e:
        return f"Flipkart automation: Searched for '{query}' but hit an issue ({e}). Page is open for you."


async def _grocery_search(page, query: str, site: str) -> str:
    """Search grocery sites (Blinkit/Zepto/BigBasket) for products."""
    urls = {
        "blinkit": f"https://blinkit.com/s/?q={query.replace(' ', '%20')}",
        "zepto": f"https://www.zeptonow.com/search?query={query.replace(' ', '%20')}",
        "bigbasket": f"https://www.bigbasket.com/ps/?q={query.replace(' ', '+')}",
    }
    url = urls.get(site, urls["blinkit"])

    try:
        await page.goto(url, timeout=20000)
        await page.wait_for_load_state("networkidle", timeout=15000)
        await page.wait_for_timeout(3000)

        # Try to click "ADD" button on the first product
        add_buttons = page.locator("button:has-text('ADD'), button:has-text('Add'), button:has-text('ADD TO CART')")
        count = await add_buttons.count()
        if count > 0:
            await add_buttons.first.click()
            await page.wait_for_timeout(2000)
            return f"Added '{query}' to your {site.title()} cart! Check the browser."

        return f"Searched for '{query}' on {site.title()}. Page is open — tap ADD on the item you want."
    except Exception as e:
        return f"{site.title()} search: Opened but hit an issue ({e}). Page is open for you."


async def _youtube_autoplay(page, query: str) -> str:
    """Open YouTube, search and play the first video."""
    try:
        await page.goto(f"https://www.youtube.com/results?search_query={query.replace(' ', '+')}", timeout=15000)
        await page.wait_for_load_state("domcontentloaded")
        await page.wait_for_timeout(2000)

        # Click first video result
        video_link = page.locator("a#video-title").first
        if await video_link.is_visible():
            await video_link.click()
            await page.wait_for_timeout(2000)
            return f"Playing '{query}' on YouTube!"

        return f"Searched for '{query}' on YouTube. Click a video to play."
    except Exception as e:
        return f"YouTube automation error: {e}"


# ─── Study & Research Tools ─────────────────────────────────────────────────

async def study_help(subject: str = "", topic: str = "", action: str = "explain") -> str:
    """Help with studying: explain topics, find resources, suggest study plans.
    Actions: explain, find_video, find_pdf, study_plan, quiz
    """
    if action == "find_video":
        # Open YouTube with educational content
        search_q = f"{subject} {topic} tutorial lecture explained".strip()
        import webbrowser
        webbrowser.open(f"https://www.youtube.com/results?search_query={search_q.replace(' ', '+')}")
        return f"Opened YouTube with study videos for {topic or subject}. I picked tutorials and lectures for you, darling."

    elif action == "find_pdf":
        search_q = f"{subject} {topic} study material notes pdf".strip()
        import webbrowser
        webbrowser.open(f"https://www.google.com/search?q={search_q.replace(' ', '+')}+filetype:pdf")
        return f"Searching for PDFs about {topic or subject}. Check your browser for the good stuff."

    elif action == "study_plan":
        return (
            f"Here's a study plan for {subject} - {topic}:\n\n"
            f"1. Start with the basics — watch a 15-20 min overview video\n"
            f"2. Read the core concepts from your textbook or notes\n"
            f"3. Practice problems — at least 10 for each subtopic\n"
            f"4. Revise with flashcards or mind maps\n"
            f"5. Take a mock test or quiz yourself\n"
            f"6. Review mistakes and weak areas\n\n"
            f"Want me to find specific videos or PDFs for any of these steps?"
        )

    elif action == "quiz":
        return (
            f"Let's quiz you on {topic or subject}! "
            f"Ask me questions about any specific topic within {subject} and I'll test your knowledge. "
            f"Or tell me the specific chapter/topic and I'll create practice questions."
        )

    else:  # explain
        # Let the LLM handle the explanation naturally
        return f"EXPLAIN_TOPIC:{subject}:{topic}"


async def research_topic(query: str, depth: str = "overview") -> str:
    """Research a topic — searches web + opens scholarly sources."""
    results = []

    # Web search for the topic
    web_result = await web_search(f"{query} research overview")
    if web_result and "No instant results" not in web_result:
        results.append(web_result)

    # Open Google Scholar in browser
    import webbrowser
    if depth == "deep":
        webbrowser.open(f"https://scholar.google.com/scholar?q={query.replace(' ', '+')}")
        results.append(f"Opened Google Scholar for deep research on '{query}'.")
    else:
        webbrowser.open(f"https://www.google.com/search?q={query.replace(' ', '+')}")
        results.append(f"Opened Google search for '{query}'.")

    if results:
        return "\n\n".join(results)
    return f"I've opened search results for '{query}'. Let me know if you need a deeper dive."


# ─── Translation Tool ───────────────────────────────────────────────────────

async def translate_text(text: str, target_lang: str = "hindi", source_lang: str = "auto") -> str:
    """Translate text between languages using free API."""
    lang_codes = {
        "hindi": "hi", "bengali": "bn", "spanish": "es", "french": "fr",
        "german": "de", "japanese": "ja", "chinese": "zh", "korean": "ko",
        "arabic": "ar", "portuguese": "pt", "russian": "ru", "italian": "it",
        "tamil": "ta", "telugu": "te", "marathi": "mr", "gujarati": "gu",
        "kannada": "kn", "malayalam": "ml", "punjabi": "pa", "urdu": "ur",
        "english": "en",
    }
    target = lang_codes.get(target_lang.lower(), target_lang.lower())
    source = "auto" if source_lang == "auto" else lang_codes.get(source_lang.lower(), source_lang.lower())

    try:
        async with httpx.AsyncClient(timeout=10.0) as client:
            resp = await client.get(
                "https://api.mymemory.translated.net/get",
                params={"q": text, "langpair": f"{source}|{target}"},
            )
            if resp.status_code == 200:
                data = resp.json()
                translated = data.get("responseData", {}).get("translatedText", "")
                if translated:
                    return f"Translation ({target_lang}): {translated}"
        return f"Could not translate. Try a different language pair."
    except Exception as e:
        return f"Translation error: {e}"


# ─── Dictionary / Define Tool ────────────────────────────────────────────────

async def define_word(word: str) -> str:
    """Look up word definition, pronunciation, and examples."""
    try:
        async with httpx.AsyncClient(timeout=8.0) as client:
            resp = await client.get(f"https://api.dictionaryapi.dev/api/v2/entries/en/{word}")
            if resp.status_code == 200:
                data = resp.json()
                if data and len(data) > 0:
                    entry = data[0]
                    word_name = entry.get("word", word)
                    phonetic = entry.get("phonetic", "")
                    meanings = entry.get("meanings", [])

                    parts = [f"{word_name} {phonetic}"]
                    for m in meanings[:3]:
                        pos = m.get("partOfSpeech", "")
                        defs = m.get("definitions", [])
                        if defs:
                            definition = defs[0].get("definition", "")
                            example = defs[0].get("example", "")
                            parts.append(f"  ({pos}) {definition}")
                            if example:
                                parts.append(f"    Example: {example}")
                    return "\n".join(parts)
            return f"No definition found for '{word}'."
    except Exception as e:
        return f"Dictionary error: {e}"


# ─── News Headlines Tool ────────────────────────────────────────────────────

async def get_news(topic: str = "top", country: str = "in") -> str:
    """Get latest news headlines."""
    try:
        # Use DuckDuckGo news
        async with httpx.AsyncClient(timeout=10.0) as client:
            q = topic if topic != "top" else "latest news India"
            resp = await client.get(
                "https://api.duckduckgo.com/",
                params={"q": q, "format": "json", "no_html": "1"},
            )
            data = resp.json()

            results = []
            for t in data.get("RelatedTopics", [])[:7]:
                if isinstance(t, dict) and t.get("Text"):
                    results.append(f"- {t['Text'][:150]}")

            if results:
                return f"News about {topic}:\n" + "\n".join(results)

            # Fallback: just open Google News
            import webbrowser
            webbrowser.open(f"https://news.google.com/search?q={topic.replace(' ', '+')}")
            return f"Opening Google News for '{topic}' in your browser."

    except Exception as e:
        return f"News error: {e}"


# ─── Calculator / Math Tool ─────────────────────────────────────────────────

async def calculate(expression: str) -> str:
    """Evaluate a math expression safely."""
    import ast
    import operator

    # Safe math operations
    ops = {
        ast.Add: operator.add, ast.Sub: operator.sub,
        ast.Mult: operator.mul, ast.Div: operator.truediv,
        ast.Pow: operator.pow, ast.Mod: operator.mod,
        ast.FloorDiv: operator.floordiv,
        ast.USub: operator.neg, ast.UAdd: operator.pos,
    }

    def _eval(node):
        if isinstance(node, ast.Constant):
            if isinstance(node.value, (int, float)):
                return node.value
            raise ValueError("Only numbers allowed")
        elif isinstance(node, ast.BinOp):
            return ops[type(node.op)](_eval(node.left), _eval(node.right))
        elif isinstance(node, ast.UnaryOp):
            return ops[type(node.op)](_eval(node.operand))
        else:
            raise ValueError(f"Unsupported: {type(node)}")

    try:
        # Clean common math language
        expr = expression.lower().strip()
        expr = expr.replace("^", "**").replace("x", "*").replace("×", "*").replace("÷", "/")
        expr = expr.replace("plus", "+").replace("minus", "-").replace("times", "*")
        expr = expr.replace("divided by", "/").replace("mod", "%").replace("power", "**")

        tree = ast.parse(expr, mode='eval')
        result = _eval(tree.body)

        if isinstance(result, float) and result == int(result):
            result = int(result)
        return f"{expression} = {result}"
    except Exception:
        return f"Could not calculate '{expression}'. Try a simpler expression like '2 + 2' or '15 * 3.5'."


# ─── Joke Tool ──────────────────────────────────────────────────────────────

async def tell_joke(category: str = "any") -> str:
    """Tell a random joke."""
    try:
        async with httpx.AsyncClient(timeout=5.0) as client:
            cat = category if category in ("programming", "misc", "pun", "dark", "spooky") else "Any"
            resp = await client.get(f"https://v2.jokeapi.dev/joke/{cat}?safe-mode")
            if resp.status_code == 200:
                data = resp.json()
                if data.get("type") == "twopart":
                    return f"{data['setup']}\n\n...{data['delivery']}"
                elif data.get("type") == "single":
                    return data.get("joke", "I forgot the punchline.")
        return "My joke database is having a moment. Ask me again!"
    except Exception:
        return "I tried to fetch a joke but the internet said no. Here's one from me: Why do programmers prefer dark mode? Because light attracts bugs."


# ─── Summarize URL/Text Tool ────────────────────────────────────────────────

async def summarize_text(text: str = "", url: str = "") -> str:
    """Summarize a long text or URL content. Returns key points."""
    if url:
        try:
            async with httpx.AsyncClient(timeout=15.0, follow_redirects=True) as client:
                resp = await client.get(url)
                if resp.status_code == 200:
                    # Extract text from HTML
                    import re
                    html = resp.text
                    # Remove scripts, styles
                    html = re.sub(r'<script[^>]*>.*?</script>', '', html, flags=re.DOTALL)
                    html = re.sub(r'<style[^>]*>.*?</style>', '', html, flags=re.DOTALL)
                    # Remove HTML tags
                    text = re.sub(r'<[^>]+>', ' ', html)
                    text = re.sub(r'\s+', ' ', text).strip()
                    text = text[:5000]  # Limit
        except Exception as e:
            return f"Couldn't fetch URL: {e}"

    if not text:
        return "No text provided to summarize."

    # Let LLM summarize
    return f"SUMMARIZE_TEXT:{text[:4000]}"


# ─── Email Draft Tool ───────────────────────────────────────────────────────

async def draft_email(to: str = "", subject: str = "", body_hint: str = "") -> str:
    """Draft an email and open Gmail compose with pre-filled fields."""
    import webbrowser
    import urllib.parse

    body = body_hint
    # Ask Ollama to draft a professional email body
    if body_hint and len(body_hint) > 5:
        from config import config
        try:
            async with httpx.AsyncClient(timeout=20.0) as client:
                resp = await client.post(
                    f"{config.ollama_url}/api/generate",
                    json={
                        "model": "jarvis:latest",
                        "prompt": f"Draft a professional email body. To: {to}. Subject: {subject}. Context: {body_hint}. Write ONLY the email body text. No subject line, no 'Subject:', no headers. Keep it concise and professional. Sign off as 'Aniket'.",
                        "stream": False,
                        "options": {"num_predict": 300, "num_gpu": -1, "num_ctx": 2048},
                    },
                )
                if resp.status_code == 200:
                    body = resp.json().get("response", body_hint).strip()
        except Exception:
            pass

    # Build Gmail compose URL
    params = {"view": "cm", "fs": "1"}
    if to:
        params["to"] = to
    if subject:
        params["su"] = subject
    if body:
        params["body"] = body

    gmail_url = f"https://mail.google.com/mail/?{urllib.parse.urlencode(params)}"
    webbrowser.open(gmail_url)
    return f"Opened Gmail compose{' to ' + to if to else ''}{' about ' + subject if subject else ''}. I've drafted the email — review and hit send!"


# ─── Volume / Brightness Control ────────────────────────────────────────────

async def system_control(action: str) -> str:
    """Control system settings: brightness, volume, wifi, bluetooth, shutdown, restart, sleep, lock."""
    import subprocess
    try:
        if action in ("shutdown", "shut down"):
            subprocess.Popen(["shutdown", "/s", "/t", "30"])
            return "Shutting down in 30 seconds. Run 'shutdown /a' to abort."
        elif action == "restart":
            subprocess.Popen(["shutdown", "/r", "/t", "30"])
            return "Restarting in 30 seconds. Run 'shutdown /a' to abort."
        elif action == "sleep":
            subprocess.Popen(["rundll32.exe", "powrprof.dll,SetSuspendState", "0,1,0"])
            return "Putting the system to sleep."
        elif action == "lock":
            subprocess.Popen(["rundll32.exe", "user32.dll,LockWorkStation"])
            return "Screen locked."
        elif action == "volume_max":
            subprocess.Popen(["powershell", "-c", "(New-Object -ComObject WScript.Shell).SendKeys([char]175)"])
            return "Volume maxed out."
        elif action == "volume_min":
            subprocess.Popen(["powershell", "-c", "for($i=0;$i-lt50;$i++){(New-Object -ComObject WScript.Shell).SendKeys([char]174)}"])
            return "Volume minimized."
        elif action == "wifi_off":
            subprocess.Popen(["netsh", "interface", "set", "interface", "Wi-Fi", "admin=disable"], shell=True)
            return "Wi-Fi disabled."
        elif action == "wifi_on":
            subprocess.Popen(["netsh", "interface", "set", "interface", "Wi-Fi", "admin=enable"], shell=True)
            return "Wi-Fi enabled."
        elif action == "empty_recycle_bin":
            subprocess.Popen(["powershell", "-c", "Clear-RecycleBin -Force -ErrorAction SilentlyContinue"])
            return "Recycle bin emptied."
        elif action == "screenshot_save":
            import mss
            from mss.tools import to_png
            with mss.mss() as sct:
                img = sct.grab(sct.monitors[1])
                path = os.path.join(os.path.expanduser("~"), "Desktop", f"screenshot_{datetime.now().strftime('%Y%m%d_%H%M%S')}.png")
                with open(path, "wb") as f:
                    f.write(to_png(img.rgb, img.size))
                return f"Screenshot saved to {path}"
        else:
            return f"Unknown system action: {action}. Try: shutdown, restart, sleep, lock, volume_max, volume_min, wifi_off, wifi_on, empty_recycle_bin, screenshot_save."
    except Exception as e:
        return f"System control error: {e}"


# ─── Clipboard Tool ─────────────────────────────────────────────────────────

async def clipboard_action(action: str = "read", text: str = "") -> str:
    """Read or write to the clipboard."""
    try:
        import subprocess
        if action == "read":
            result = subprocess.run(["powershell", "-c", "Get-Clipboard"], capture_output=True, text=True, timeout=5)
            content = result.stdout.strip()
            return f"Clipboard contents: {content}" if content else "Clipboard is empty."
        elif action == "write":
            subprocess.run(["powershell", "-c", f"Set-Clipboard -Value '{text}'"], timeout=5)
            return f"Copied to clipboard: {text[:100]}"
        elif action == "clear":
            subprocess.run(["powershell", "-c", "Set-Clipboard -Value $null"], timeout=5)
            return "Clipboard cleared."
    except Exception as e:
        return f"Clipboard error: {e}"


# ─── Speedtest Tool ─────────────────────────────────────────────────────────

async def internet_speed() -> str:
    """Quick internet speed check using a fast download test."""
    import time
    try:
        async with httpx.AsyncClient(timeout=20.0) as client:
            # Download a known file and measure speed
            url = "https://speed.cloudflare.com/__down?bytes=5000000"  # 5MB
            start = time.time()
            resp = await client.get(url)
            elapsed = time.time() - start
            if resp.status_code == 200:
                size_mb = len(resp.content) / (1024 * 1024)
                speed_mbps = (size_mb * 8) / elapsed
                return f"Download speed: {speed_mbps:.1f} Mbps (tested {size_mb:.1f} MB in {elapsed:.1f}s)"
        return "Speed test failed. Check your connection."
    except Exception as e:
        return f"Speed test error: {e}"


# ─── QR Code Generator ──────────────────────────────────────────────────────

async def generate_qr(data: str) -> str:
    """Generate a QR code and save to Desktop."""
    try:
        # Use Google Charts API for QR generation (no dependency needed)
        import webbrowser
        url = f"https://api.qrserver.com/v1/create-qr-code/?size=300x300&data={data.replace(' ', '%20')}"
        webbrowser.open(url)
        return f"QR code for '{data[:50]}' opened in browser."
    except Exception as e:
        return f"QR code error: {e}"


# ─── IP / Location Info ─────────────────────────────────────────────────────

async def ip_info() -> str:
    """Get public IP address and location info."""
    try:
        async with httpx.AsyncClient(timeout=8.0) as client:
            resp = await client.get("https://ipapi.co/json/")
            if resp.status_code == 200:
                d = resp.json()
                return (
                    f"Public IP: {d.get('ip', 'unknown')}\n"
                    f"Location: {d.get('city', '?')}, {d.get('region', '?')}, {d.get('country_name', '?')}\n"
                    f"ISP: {d.get('org', 'unknown')}\n"
                    f"Timezone: {d.get('timezone', 'unknown')}"
                )
        return "Could not fetch IP info."
    except Exception as e:
        return f"IP info error: {e}"


# ─── Crypto / Stock Price Tool ───────────────────────────────────────────────

async def get_price(symbol: str) -> str:
    """Get cryptocurrency or stock price."""
    symbol_upper = symbol.upper().strip()
    try:
        # Try crypto first (CoinGecko)
        crypto_map = {
            "BTC": "bitcoin", "ETH": "ethereum", "SOL": "solana",
            "DOGE": "dogecoin", "ADA": "cardano", "XRP": "ripple",
            "DOT": "polkadot", "MATIC": "matic-network", "AVAX": "avalanche-2",
            "BITCOIN": "bitcoin", "ETHEREUM": "ethereum",
        }
        coin_id = crypto_map.get(symbol_upper)
        if coin_id:
            async with httpx.AsyncClient(timeout=8.0) as client:
                resp = await client.get(
                    f"https://api.coingecko.com/api/v3/simple/price?ids={coin_id}&vs_currencies=usd,inr&include_24hr_change=true"
                )
                if resp.status_code == 200:
                    data = resp.json().get(coin_id, {})
                    if data:
                        usd = data.get("usd", 0)
                        inr = data.get("inr", 0)
                        change = data.get("usd_24h_change", 0)
                        arrow = "up" if change >= 0 else "down"
                        return f"{symbol_upper}: ${usd:,.2f} (₹{inr:,.0f}) — {arrow} {abs(change):.2f}% in 24h"

        # Otherwise try as stock (Yahoo Finance scrape would need API key, so use web search)
        result = await web_search(f"{symbol_upper} stock price today")
        return result if result else f"Could not find price for {symbol_upper}."
    except Exception as e:
        return f"Price lookup error: {e}"


# ─── Timer / Stopwatch Tool ─────────────────────────────────────────────────

_active_timers: dict[str, datetime] = {}


async def timer_action(action: str = "start", name: str = "default", duration: str = "") -> str:
    """Start/stop/check timers and stopwatches."""
    if action == "start":
        _active_timers[name] = datetime.now()
        return f"Timer '{name}' started."
    elif action == "stop" or action == "check":
        if name in _active_timers:
            elapsed = datetime.now() - _active_timers[name]
            mins = int(elapsed.total_seconds() // 60)
            secs = int(elapsed.total_seconds() % 60)
            if action == "stop":
                del _active_timers[name]
            return f"Timer '{name}': {mins}m {secs}s elapsed."
        return f"No timer named '{name}' is running."
    elif action == "list":
        if not _active_timers:
            return "No timers running."
        lines = []
        for n, start in _active_timers.items():
            elapsed = datetime.now() - start
            lines.append(f"  {n}: {int(elapsed.total_seconds())}s")
        return "Active timers:\n" + "\n".join(lines)
    return f"Unknown timer action: {action}"


# ─── Create Folder Tool ─────────────────────────────────────────────────────

async def create_folder(path: str = "", name: str = "") -> str:
    """Create a new folder on the device."""
    if not name and not path:
        return "Please specify a folder name."

    if path and not name:
        target = Path(path)
    elif name and not path:
        # Default: create on Desktop
        target = Path(os.path.expanduser("~")) / "Desktop" / name
    else:
        target = Path(path) / name

    try:
        target.mkdir(parents=True, exist_ok=True)
        return f"Created folder: {target}"
    except Exception as e:
        return f"Could not create folder: {e}"


# ─── Write Code Tool ────────────────────────────────────────────────────────

async def write_code(filename: str = "main.py", code: str = "", folder: str = "", language: str = "") -> str:
    """Create a code file with content and open it in VS Code."""
    # If no code provided, generate it using Ollama
    if not code and language:
        from config import config
        try:
            # Detect appropriate filename/extension from language hint
            lang_lower = language.lower()
            ext_map = {
                "python": ".py", "javascript": ".js", "typescript": ".ts",
                "java": ".java", "html": ".html", "css": ".css",
                "cpp": ".cpp", "c++": ".cpp", "c": ".c", "rust": ".rs",
                "go": ".go", "ruby": ".rb", "php": ".php", "swift": ".swift",
            }
            for lang_key, ext in ext_map.items():
                if lang_key in lang_lower:
                    if filename == "main.py":  # default, change it
                        filename = f"main{ext}"
                    break

            async with httpx.AsyncClient(timeout=30.0) as client:
                resp = await client.post(
                    f"{config.ollama_url}/api/generate",
                    json={
                        "model": "jarvis:latest",
                        "prompt": f"Write clean, working code for the following request. Output ONLY the code with proper syntax, no explanations, no markdown, no backticks. Request: {language}",
                        "stream": False,
                        "options": {"num_predict": 800, "temperature": 0.3, "num_gpu": -1, "num_ctx": 2048},
                    },
                )
                if resp.status_code == 200:
                    code = resp.json().get("response", "").strip()
                    # Clean up any markdown code fences
                    if code.startswith("```"):
                        lines = code.split("\n")
                        lines = [l for l in lines if not l.startswith("```")]
                        code = "\n".join(lines)
        except Exception:
            pass

    if not code:
        return "I need to know what code to write. Try: 'write a python script for a calculator'"

    # Determine folder
    if folder:
        target_dir = Path(folder)
    else:
        target_dir = Path(os.path.expanduser("~")) / "Desktop"

    target_dir.mkdir(parents=True, exist_ok=True)
    file_path = target_dir / filename

    try:
        file_path.write_text(code, encoding="utf-8")

        # Open in VS Code
        try:
            subprocess.Popen(["code", str(file_path)], shell=True, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
        except Exception:
            # Fallback: open in notepad
            subprocess.Popen(["notepad", str(file_path)], stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)

        return f"Created {filename} and opened it in VS Code. File at: {file_path}"
    except Exception as e:
        return f"Error writing code file: {e}"


# ─── Write Document Tool ────────────────────────────────────────────────────

async def write_document(text: str = "", filename: str = "", app: str = "notepad") -> str:
    """Write text to a file and open it in Notepad or Word."""
    import asyncio

    # If text looks like a request rather than content, generate content with Ollama
    request_words = ["write", "type", "create", "make", "draft", "note", "about", "explain", "describe", "essay", "letter", "report", "article"]
    if text and any(w in text.lower()[:80] for w in request_words):
        from config import config
        try:
            async with httpx.AsyncClient(timeout=45.0) as client:
                resp = await client.post(
                    f"{config.ollama_url}/api/generate",
                    json={
                        "model": "jarvis:latest",
                        "prompt": f"Write the following content as a well-formatted document. No markdown, plain text only. Be detailed, thorough and professional. Write at least 500 words. Request: {text}",
                        "stream": False,
                        "options": {"num_predict": 1200, "temperature": 0.6, "num_gpu": -1, "num_ctx": 2048},
                    },
                )
                if resp.status_code == 200:
                    generated = resp.json().get("response", "").strip()
                    if generated and len(generated) > 30:
                        text = generated
        except Exception:
            pass

    if not text:
        return "No text provided to write."

    app_lower = app.lower().strip()

    # Determine filename and extension
    if not filename:
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        if app_lower == "word":
            filename = f"jarvis_doc_{timestamp}.docx"
        else:
            filename = f"jarvis_doc_{timestamp}.txt"

    file_path = Path(os.path.expanduser("~")) / "Desktop" / filename

    try:
        if app_lower == "word" and filename.endswith(".docx"):
            # Try python-docx if available
            try:
                from docx import Document
                doc = Document()
                # Add title
                doc.add_heading("JARVIS Document", level=1)
                # Add paragraphs
                for paragraph in text.split("\n"):
                    if paragraph.strip():
                        doc.add_paragraph(paragraph.strip())
                doc.save(str(file_path))
            except ImportError:
                # Fallback: write as .txt and open in Word
                file_path = file_path.with_suffix(".txt")
                file_path.write_text(text, encoding="utf-8")

            await asyncio.sleep(0.3)

            # Open in Word explicitly via winword.exe
            try:
                # Try finding Word
                winword_paths = [
                    r"C:\Program Files\Microsoft Office\root\Office16\WINWORD.EXE",
                    r"C:\Program Files (x86)\Microsoft Office\root\Office16\WINWORD.EXE",
                    r"C:\Program Files\Microsoft Office\Office16\WINWORD.EXE",
                    r"C:\Program Files (x86)\Microsoft Office\Office16\WINWORD.EXE",
                ]
                word_found = False
                for wp in winword_paths:
                    if os.path.exists(wp):
                        subprocess.Popen([wp, str(file_path)], stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
                        word_found = True
                        break
                if not word_found:
                    # Fallback to os.startfile or shell start
                    try:
                        os.startfile(str(file_path))  # type: ignore[attr-defined]
                    except Exception:
                        subprocess.Popen(["cmd", "/c", "start", "", str(file_path)], stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
            except Exception:
                subprocess.Popen(["notepad", str(file_path)], stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
            return f"Created document and opened in Word. File at: {file_path}"
        else:
            # Notepad
            file_path.write_text(text, encoding="utf-8")
            await asyncio.sleep(0.3)
            subprocess.Popen(["notepad", str(file_path)], stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
            return f"Created document and opened in Notepad. File at: {file_path}"
    except Exception as e:
        return f"Error writing document: {e}"


# ─── Draw Diagram Tool (Pillow + Paint) ─────────────────────────────────────

def _draw_house(draw, w=800, h=600):
    """Draw a simple house."""
    # Main body
    draw.rectangle([200, 250, 600, 500], outline="black", width=3)
    # Roof (triangle)
    draw.polygon([(180, 250), (400, 80), (620, 250)], outline="black", width=3)
    # Door
    draw.rectangle([350, 380, 450, 500], outline="brown", fill="brown", width=2)
    # Door knob
    draw.ellipse([430, 435, 440, 445], fill="gold")
    # Windows
    draw.rectangle([240, 300, 320, 370], outline="blue", fill="lightblue", width=2)
    draw.line([(280, 300), (280, 370)], fill="blue", width=2)
    draw.line([(240, 335), (320, 335)], fill="blue", width=2)
    draw.rectangle([480, 300, 560, 370], outline="blue", fill="lightblue", width=2)
    draw.line([(520, 300), (520, 370)], fill="blue", width=2)
    draw.line([(480, 335), (560, 335)], fill="blue", width=2)
    # Chimney
    draw.rectangle([500, 100, 540, 200], outline="red", fill="red", width=2)
    # Ground
    draw.line([(0, 500), (800, 500)], fill="green", width=3)
    # Labels
    draw.text((340, 540), "House", fill="black")


def _draw_neural_network(draw, w=800, h=600):
    """Draw a simple neural network diagram."""
    layers = [[150, [150, 250, 350, 450]],
              [350, [120, 220, 320, 420, 480]],
              [550, [200, 300, 400]],
              [700, [280, 350]]]

    node_r = 18
    colors = ["#4CAF50", "#2196F3", "#FF9800", "#F44336"]

    # Draw connections first
    for i in range(len(layers) - 1):
        x1, ys1 = layers[i]
        x2, ys2 = layers[i + 1]
        for y1 in ys1:
            for y2 in ys2:
                draw.line([(x1, y1), (x2, y2)], fill="#CCCCCC", width=1)

    # Draw nodes
    for i, (x, ys) in enumerate(layers):
        for y in ys:
            draw.ellipse([x - node_r, y - node_r, x + node_r, y + node_r],
                        fill=colors[i], outline="black", width=2)

    # Labels
    labels = ["Input\nLayer", "Hidden\nLayer 1", "Hidden\nLayer 2", "Output\nLayer"]
    for i, (x, ys) in enumerate(layers):
        draw.text((x - 25, 530), labels[i], fill="black")
    draw.text((280, 30), "Neural Network Architecture", fill="black")


def _draw_flowchart(draw, desc="", w=800, h=600):
    """Draw a basic flowchart."""
    # Start oval
    draw.ellipse([300, 20, 500, 70], outline="green", fill="lightgreen", width=2)
    draw.text((370, 35), "Start", fill="black")
    # Arrow
    draw.line([(400, 70), (400, 110)], fill="black", width=2)
    draw.polygon([(395, 105), (405, 105), (400, 115)], fill="black")
    # Process box
    draw.rectangle([280, 115, 520, 170], outline="blue", fill="lightblue", width=2)
    draw.text((330, 132), "Process / Input", fill="black")
    # Arrow
    draw.line([(400, 170), (400, 210)], fill="black", width=2)
    draw.polygon([(395, 205), (405, 205), (400, 215)], fill="black")
    # Decision diamond
    draw.polygon([(400, 215), (520, 290), (400, 365), (280, 290)], outline="orange", fill="lightyellow", width=2)
    draw.text((355, 278), "Decision?", fill="black")
    # Yes branch
    draw.line([(520, 290), (620, 290)], fill="black", width=2)
    draw.text((540, 270), "Yes", fill="green")
    draw.rectangle([620, 260, 760, 320], outline="blue", fill="lightblue", width=2)
    draw.text((650, 280), "Action A", fill="black")
    # No branch
    draw.line([(400, 365), (400, 410)], fill="black", width=2)
    draw.text((410, 370), "No", fill="red")
    draw.rectangle([280, 410, 520, 465], outline="blue", fill="lightblue", width=2)
    draw.text((350, 428), "Action B", fill="black")
    # End
    draw.line([(400, 465), (400, 510)], fill="black", width=2)
    draw.polygon([(395, 505), (405, 505), (400, 515)], fill="black")
    draw.ellipse([300, 515, 500, 565], outline="red", fill="#FFCCCC", width=2)
    draw.text((380, 530), "End", fill="black")


def _draw_tree(draw, w=800, h=600):
    """Draw a simple tree."""
    # Trunk
    draw.rectangle([370, 350, 430, 520], fill="brown", outline="brown", width=2)
    # Canopy (overlapping circles for foliage)
    for cx, cy, r in [(400, 280, 80), (340, 300, 60), (460, 300, 60), (370, 230, 55), (430, 230, 55), (400, 200, 50)]:
        draw.ellipse([cx - r, cy - r, cx + r, cy + r], fill="green", outline="darkgreen", width=2)
    # Ground
    draw.line([(0, 520), (800, 520)], fill="green", width=3)
    draw.text((370, 550), "Tree", fill="black")


def _draw_generic(draw, subject, w=800, h=600):
    """Draw a generic labeled diagram."""
    draw.rectangle([50, 50, 750, 550], outline="black", width=2)
    draw.text((300, 20), f"Diagram: {subject}", fill="black")
    # Central shape
    draw.ellipse([250, 200, 550, 400], outline="blue", fill="lightblue", width=3)
    draw.text((340, 285), subject[:20], fill="black")
    # Decorative connecting elements
    for i, angle_y in enumerate([100, 200, 300, 400, 500]):
        draw.line([(100, angle_y), (250, 300)], fill="gray", width=1)
        draw.ellipse([70, angle_y - 15, 110, angle_y + 15], fill="orange", outline="black")
        draw.line([(550, 300), (700, angle_y)], fill="gray", width=1)
        draw.ellipse([690, angle_y - 15, 730, angle_y + 15], fill="green", outline="black")


async def draw_diagram(subject: str, description: str = "") -> str:
    """Draw a diagram/illustration using Pillow and open it in MS Paint."""
    try:
        from PIL import Image, ImageDraw, ImageFont
    except ImportError:
        return "Drawing requires Pillow. Install with: pip install Pillow"

    img = Image.new('RGB', (800, 600), 'white')
    draw = ImageDraw.Draw(img)

    # Try to use a readable font
    try:
        font = ImageFont.truetype("arial.ttf", 16)
        font_title = ImageFont.truetype("arial.ttf", 20)
    except Exception:
        font = ImageFont.load_default()
        font_title = font

    subject_lower = subject.lower()
    if 'house' in subject_lower or 'home' in subject_lower or 'building' in subject_lower:
        _draw_house(draw)
    elif 'tree' in subject_lower or 'plant' in subject_lower:
        _draw_tree(draw)
    elif 'neural' in subject_lower or 'network' in subject_lower or 'nn' in subject_lower:
        _draw_neural_network(draw)
    elif 'flowchart' in subject_lower or 'flow' in subject_lower or 'process' in subject_lower:
        _draw_flowchart(draw, description)
    else:
        _draw_generic(draw, subject)

    # Save to Desktop — ensure flush
    safe_name = subject.replace(' ', '_').replace('/', '_')[:30]
    path = os.path.join(os.path.expanduser("~"), "Desktop", f"jarvis_drawing_{safe_name}.png")
    img.save(path, "PNG")
    img.close()

    # Small delay to ensure file is written to disk
    import asyncio
    await asyncio.sleep(0.5)

    # Open in Paint using os.startfile for reliability on Windows
    try:
        os.startfile(path)  # type: ignore[attr-defined]
    except Exception:
        subprocess.Popen(["mspaint", path], stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
    return f"Drew a {subject} diagram and opened it in Paint! File saved at: {path}"


# ─── Gmail Check Tool ───────────────────────────────────────────────────────

async def check_gmail() -> str:
    """Check for unread emails in Gmail via IMAP."""
    import imaplib
    import email as email_lib
    from email.header import decode_header

    gmail_user = os.environ.get("JARVIS_EMAIL_USER", "aniketsupermails2005@gmail.com")
    gmail_pass = os.environ.get("JARVIS_EMAIL_PASSWORD", "")

    if not gmail_pass:
        # Fallback: just open Gmail in browser
        import webbrowser
        webbrowser.open("https://mail.google.com")
        return "Opened Gmail in your browser. To enable email checking, set JARVIS_EMAIL_PASSWORD in your .env file with a Gmail App Password."

    try:
        mail = imaplib.IMAP4_SSL("imap.gmail.com")
        mail.login(gmail_user, gmail_pass)
        mail.select("inbox")

        # Search for unread emails (last 5)
        status, messages = mail.search(None, "UNSEEN")
        mail_ids = messages[0].split()

        if not mail_ids:
            mail.logout()
            return "No new unread emails. Your inbox is clean!"

        results = []
        # Get last 5 unread
        for mid in mail_ids[-5:]:
            status, msg_data = mail.fetch(mid, "(RFC822)")
            if status != "OK":
                continue
            msg = email_lib.message_from_bytes(msg_data[0][1])

            # Decode subject
            subject_raw = msg.get("Subject", "No Subject")
            decoded_parts = decode_header(subject_raw)
            subject = ""
            for part, enc in decoded_parts:
                if isinstance(part, bytes):
                    subject += part.decode(enc or "utf-8", errors="replace")
                else:
                    subject += str(part)

            sender = msg.get("From", "Unknown")
            date = msg.get("Date", "")[:25]
            results.append(f"• From: {sender}\n  Subject: {subject}\n  Date: {date}")

        mail.logout()
        count = len(mail_ids)
        header = f"You have {count} unread email{'s' if count != 1 else ''}:\n\n"
        return header + "\n\n".join(results)

    except imaplib.IMAP4.error as e:
        if "AUTHENTICATIONFAILED" in str(e):
            return "Gmail authentication failed. Make sure JARVIS_EMAIL_PASSWORD is a valid Gmail App Password (not your regular password)."
        return f"Gmail error: {e}"
    except Exception as e:
        return f"Email check error: {e}"


# ─── WhatsApp Check Tool ────────────────────────────────────────────────────

async def check_whatsapp() -> str:
    """Open WhatsApp Web to check for new messages."""
    try:
        _ensure_playwright_browsers()
        from playwright.async_api import async_playwright

        pw = await async_playwright().start()
        browser = await pw.chromium.launch(
            headless=False,
            args=["--start-maximized"],
            channel="msedge"  # Use Edge to keep WhatsApp session
        )
        _pw_refs.append(pw)

        # Use persistent context to keep WhatsApp logged in
        user_data = os.path.join(os.environ.get("LOCALAPPDATA", ""), "JARVIS", "whatsapp_profile")
        context = await browser.new_context(
            viewport={"width": 1366, "height": 900},
        )
        page = await context.new_page()

        await page.goto("https://web.whatsapp.com", timeout=30000)
        await page.wait_for_timeout(5000)

        # Check if we need to scan QR code
        qr = page.locator("canvas[aria-label='Scan this QR code to link a device!']")
        if await qr.is_visible(timeout=3000):
            return "WhatsApp Web opened. Please scan the QR code with your phone to connect. I'll keep the browser open for you."

        # Already logged in — check for unread chats
        await page.wait_for_timeout(3000)
        unread_badges = page.locator("span[aria-label*='unread message']")
        count = await unread_badges.count()

        if count > 0:
            return f"WhatsApp Web is open. You have messages from {count} chat{'s' if count > 1 else ''} with unread messages. Check your browser!"
        return "WhatsApp Web is open. No new unread messages right now."

    except ImportError:
        import webbrowser
        webbrowser.open("https://web.whatsapp.com")
        return "Opened WhatsApp Web in your browser."
    except Exception as e:
        import webbrowser
        webbrowser.open("https://web.whatsapp.com")
        return f"Opened WhatsApp Web in your browser. (Automation note: {e})"


# ─── Git Operations ─────────────────────────────────────────────────────────

async def git_commit(message: str = "", path: str = "") -> str:
    """Stage all changes and commit with a message. Auto-generates message if none given."""
    cwd = path or str(Path.home())
    try:
        # Check if git repo
        check = subprocess.run(["git", "rev-parse", "--is-inside-work-tree"],
                               capture_output=True, text=True, cwd=cwd, timeout=5)
        if check.returncode != 0:
            return f"Not a git repository: {cwd}"

        # Get status
        status = subprocess.run(["git", "status", "--porcelain"],
                                capture_output=True, text=True, cwd=cwd, timeout=10)
        changes = status.stdout.strip()
        if not changes:
            return "No changes to commit — working tree is clean."

        # Auto-generate message if not provided
        if not message:
            files_changed = [line[3:] for line in changes.split("\n") if line.strip()]
            file_list = ", ".join(files_changed[:5])
            if len(files_changed) > 5:
                file_list += f" (+{len(files_changed)-5} more)"
            message = f"Update {file_list}"

        # Stage all
        subprocess.run(["git", "add", "-A"], cwd=cwd, timeout=10,
                        capture_output=True, text=True)

        # Commit
        result = subprocess.run(["git", "commit", "-m", message],
                                capture_output=True, text=True, cwd=cwd, timeout=15)
        output = result.stdout.strip() or result.stderr.strip()
        return f"Git commit: {output}"
    except subprocess.TimeoutExpired:
        return "Git operation timed out."
    except Exception as e:
        return f"Git error: {e}"


async def git_status(path: str = "") -> str:
    """Get current git status including branch, changes, and recent commits."""
    cwd = path or str(Path.home())
    try:
        lines = []
        # Branch
        branch = subprocess.run(["git", "branch", "--show-current"],
                                capture_output=True, text=True, cwd=cwd, timeout=5)
        lines.append(f"Branch: {branch.stdout.strip()}")

        # Status
        status = subprocess.run(["git", "status", "--short"],
                                capture_output=True, text=True, cwd=cwd, timeout=5)
        changes = status.stdout.strip()
        if changes:
            lines.append(f"Changes:\n{changes}")
        else:
            lines.append("Working tree clean.")

        # Recent commits
        log = subprocess.run(["git", "log", "--oneline", "-5"],
                             capture_output=True, text=True, cwd=cwd, timeout=5)
        if log.stdout.strip():
            lines.append(f"Recent commits:\n{log.stdout.strip()}")

        return "\n".join(lines)
    except Exception as e:
        return f"Git status error: {e}"


async def git_push(path: str = "") -> str:
    """Push committed changes to remote."""
    cwd = path or str(Path.home())
    try:
        result = subprocess.run(["git", "push"],
                                capture_output=True, text=True, cwd=cwd, timeout=30)
        output = result.stdout.strip() or result.stderr.strip()
        if result.returncode == 0:
            return f"Pushed successfully: {output}"
        return f"Push failed: {output}"
    except Exception as e:
        return f"Git push error: {e}"


async def git_pull(path: str = "") -> str:
    """Pull latest changes from remote."""
    cwd = path or str(Path.home())
    try:
        result = subprocess.run(["git", "pull"],
                                capture_output=True, text=True, cwd=cwd, timeout=30)
        return result.stdout.strip() or result.stderr.strip() or "Already up to date."
    except Exception as e:
        return f"Git pull error: {e}"


# ─── Run Python Script ──────────────────────────────────────────────────────

async def run_python(code: str = "", file: str = "") -> str:
    """Execute Python code or run a Python file. Returns stdout."""
    try:
        if file:
            p = Path(file).expanduser().resolve()
            if not p.exists():
                return f"File not found: {file}"
            result = subprocess.run(
                ["python", str(p)],
                capture_output=True, text=True, timeout=30,
                cwd=str(p.parent)
            )
        elif code:
            result = subprocess.run(
                ["python", "-c", code],
                capture_output=True, text=True, timeout=30,
                cwd=str(Path.home())
            )
        else:
            return "Provide either 'code' or 'file' to run."

        output = result.stdout[:3000] if result.stdout else ""
        error = result.stderr[:1000] if result.stderr else ""
        if result.returncode != 0:
            return f"Error (exit {result.returncode}):\n{error}\n{output}"
        return output or "(no output)"
    except subprocess.TimeoutExpired:
        return "Script timed out (30s limit)."
    except Exception as e:
        return f"Python execution error: {e}"


# ─── Create Project Scaffold ────────────────────────────────────────────────

async def create_project(name: str, template: str = "python") -> str:
    """Create a new project folder with common boilerplate.
    Templates: python, web, node, react, flask, fastapi
    """
    base = Path.home() / "Desktop" / name
    base.mkdir(parents=True, exist_ok=True)

    if template in ("python", "py"):
        (base / "main.py").write_text('"""Main entry point."""\n\n\ndef main():\n    print("Hello from ' + name + '!")\n\n\nif __name__ == "__main__":\n    main()\n')
        (base / "requirements.txt").write_text("# Add your dependencies here\n")
        (base / "README.md").write_text(f"# {name}\n\nA Python project.\n")
        (base / ".gitignore").write_text("__pycache__/\n*.pyc\n.venv/\nvenv/\n.env\ndist/\n*.egg-info/\n")

    elif template in ("web", "html"):
        (base / "index.html").write_text(f'<!DOCTYPE html>\n<html lang="en">\n<head>\n  <meta charset="UTF-8">\n  <meta name="viewport" content="width=device-width, initial-scale=1.0">\n  <title>{name}</title>\n  <link rel="stylesheet" href="style.css">\n</head>\n<body>\n  <h1>{name}</h1>\n  <script src="script.js"></script>\n</body>\n</html>\n')
        (base / "style.css").write_text("* { margin: 0; padding: 0; box-sizing: border-box; }\nbody { font-family: system-ui; padding: 2rem; }\n")
        (base / "script.js").write_text(f'console.log("{name} loaded");\n')

    elif template in ("node", "nodejs"):
        (base / "index.js").write_text(f'console.log("{name} started");\n')
        (base / "package.json").write_text(json.dumps({"name": name, "version": "1.0.0", "main": "index.js", "scripts": {"start": "node index.js"}}, indent=2))
        (base / ".gitignore").write_text("node_modules/\n.env\ndist/\n")

    elif template in ("fastapi", "api"):
        (base / "main.py").write_text(f'from fastapi import FastAPI\n\napp = FastAPI(title="{name}")\n\n@app.get("/")\ndef root():\n    return {{"message": "Hello from {name}!"}}\n')
        (base / "requirements.txt").write_text("fastapi\nuvicorn[standard]\n")
        (base / ".gitignore").write_text("__pycache__/\n*.pyc\n.venv/\nvenv/\n.env\n")

    else:
        (base / "README.md").write_text(f"# {name}\n\nProject created by JARVIS.\n")

    # Init git
    subprocess.run(["git", "init"], cwd=str(base), capture_output=True, timeout=5)

    # Open in VS Code
    try:
        subprocess.Popen(["code", str(base)], shell=True, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
    except Exception:
        pass

    return f"Created {template} project '{name}' at {base} and opened in VS Code."


# Tool registry
TOOLS = {
    "web_search": {"fn": web_search, "desc": "Search the web for information", "args": ["query"]},
    "weather": {"fn": get_weather, "desc": "Get current weather for a city", "args": ["city"]},
    "terminal": {"fn": run_terminal_command, "desc": "Run a terminal/shell command", "args": ["command"]},
    "screenshot": {"fn": capture_screen, "desc": "Capture a screenshot of the screen", "args": []},
    "system_info": {"fn": get_system_info, "desc": "Get CPU, RAM, disk stats", "args": []},
    "read_file": {"fn": read_file, "desc": "Read a file's contents", "args": ["filepath"]},
    "write_file": {"fn": write_file, "desc": "Write content to a file", "args": ["filepath", "content"]},
    "datetime": {"fn": get_datetime, "desc": "Get current date and time", "args": []},
    "vision_analyze": {"fn": vision_analyze, "desc": "Capture screen and analyze with YOLO + AI vision", "args": ["query"]},
    "scheduler": {"fn": scheduler, "desc": "Manage reminders: add/list/remove/clear", "args": ["action", "text", "time"]},
    "media_control": {"fn": media_control, "desc": "Control media playback: play_pause/next/previous/volume_up/volume_down/mute", "args": ["action"]},
    "mobile_sync": {"fn": mobile_sync, "desc": "Send notifications/emails/whatsapp or view profile", "args": ["action", "message", "to"]},
    "open_url": {"fn": open_url, "desc": "Open a URL in the default browser", "args": ["url"]},
    "open_app": {"fn": open_app, "desc": "Open a website: youtube, amazon, flipkart, spotify, etc.", "args": ["name"]},
    "youtube_play": {"fn": youtube_play, "desc": "Search and play a video/song on YouTube", "args": ["query"]},
    "search_files": {"fn": search_files, "desc": "Search for files/folders on the device by name or pattern", "args": ["query", "path"]},
    "launch_app": {"fn": launch_app, "desc": "Find and launch a desktop application (notepad, chrome, vscode, etc.)", "args": ["name"]},
    "browser_action": {"fn": browser_action, "desc": "Automate browser: search, add to cart, open sites, research, find PDFs", "args": ["action", "query", "site"]},
    "study_help": {"fn": study_help, "desc": "Study assistance: explain topics, find videos/PDFs, create study plans, quiz", "args": ["subject", "topic", "action"]},
    "research_topic": {"fn": research_topic, "desc": "Research a topic with web search + scholarly sources", "args": ["query", "depth"]},
    "translate": {"fn": translate_text, "desc": "Translate text to another language", "args": ["text", "target_lang"]},
    "define": {"fn": define_word, "desc": "Look up word definition and meaning", "args": ["word"]},
    "news": {"fn": get_news, "desc": "Get latest news headlines on a topic", "args": ["topic"]},
    "calculate": {"fn": calculate, "desc": "Evaluate a math expression", "args": ["expression"]},
    "joke": {"fn": tell_joke, "desc": "Tell a random joke", "args": ["category"]},
    "summarize": {"fn": summarize_text, "desc": "Summarize text or a URL", "args": ["text", "url"]},
    "draft_email": {"fn": draft_email, "desc": "Draft a professional email", "args": ["to", "subject", "body_hint"]},
    "system_control": {"fn": system_control, "desc": "System control: shutdown, restart, sleep, lock, wifi, volume, recycle bin", "args": ["action"]},
    "clipboard": {"fn": clipboard_action, "desc": "Read, write, or clear clipboard", "args": ["action", "text"]},
    "speed_test": {"fn": internet_speed, "desc": "Test internet download speed", "args": []},
    "qr_code": {"fn": generate_qr, "desc": "Generate a QR code for text/URL", "args": ["data"]},
    "ip_info": {"fn": ip_info, "desc": "Get public IP address and location", "args": []},
    "price": {"fn": get_price, "desc": "Get crypto or stock price", "args": ["symbol"]},
    "timer": {"fn": timer_action, "desc": "Start/stop/check timers and stopwatches", "args": ["action", "name"]},
    "create_folder": {"fn": create_folder, "desc": "Create a new folder on the device", "args": ["path", "name"]},
    "write_code": {"fn": write_code, "desc": "Create a code file and open in VS Code", "args": ["filename", "code", "folder", "language"]},
    "write_document": {"fn": write_document, "desc": "Write text to a file, open in Notepad or Word", "args": ["text", "filename", "app"]},
    "draw_diagram": {"fn": draw_diagram, "desc": "Draw a diagram/illustration in MS Paint", "args": ["subject", "description"]},
    "check_gmail": {"fn": check_gmail, "desc": "Check for unread emails in Gmail", "args": []},
    "check_whatsapp": {"fn": check_whatsapp, "desc": "Open WhatsApp Web to check messages", "args": []},
    "git_commit": {"fn": git_commit, "desc": "Stage and commit git changes", "args": ["message", "path"]},
    "git_status": {"fn": git_status, "desc": "Show git branch, changes, and recent commits", "args": ["path"]},
    "git_push": {"fn": git_push, "desc": "Push commits to remote", "args": ["path"]},
    "git_pull": {"fn": git_pull, "desc": "Pull latest changes from remote", "args": ["path"]},
    "run_python": {"fn": run_python, "desc": "Execute Python code or run a .py file", "args": ["code", "file"]},
    "create_project": {"fn": create_project, "desc": "Create a new project folder with boilerplate", "args": ["name", "template"]},
}


def get_tools_description() -> str:
    """Get a description of all available tools for the system prompt."""
    lines = []
    for name, info in TOOLS.items():
        args = ", ".join(info["args"]) if info["args"] else "none"
        lines.append(f"  - {name}({args}): {info['desc']}")
    return "\n".join(lines)


async def execute_tool(name: str, args: dict) -> str:
    """Execute a tool by name with given args."""
    tool = TOOLS.get(name)
    if not tool:
        return f"Unknown tool: {name}"
    try:
        return await tool["fn"](**args)
    except Exception as e:
        return f"Tool error ({name}): {e}"
